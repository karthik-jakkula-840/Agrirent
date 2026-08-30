import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        if level == 1:
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(22, 101, 52) # Dark Green
        elif level == 2:
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 64, 175) # Dark Blue
        elif level == 3:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(55, 65, 81) # Gray-700
    return h

def add_bullet_point(doc, title, desc):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run_t = p.add_run(f"{title}: ")
    run_t.bold = True
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(17, 24, 39)
    
    run_d = p.add_run(desc)
    run_d.font.name = 'Calibri'
    run_d.font.size = Pt(10.5)
    run_d.font.color.rgb = RGBColor(75, 85, 99)
    return p

def add_image_safe(doc, image_path, caption="", width=Inches(6.0)):
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(image_path, width=width)
        
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run(f"Figure: {caption}")
            run_cap.font.name = 'Calibri'
            run_cap.font.size = Pt(9.5)
            run_cap.font.italic = True
            run_cap.font.color.rgb = RGBColor(107, 114, 128)
    else:
        print(f"Warning: Image not found at {image_path}")

def generate_document():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Set base font
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(31, 41, 55)

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    img_dir = os.path.join(base_dir, 'doc_screenshots')

    # ==========================================
    # 1. TITLE / COVER PAGE
    # ==========================================
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(36)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AGRIRENT PLATFORM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(22, 101, 52)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("Next-Generation Agricultural Equipment Rental Marketplace\nComprehensive Project & Feature Documentation")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(75, 85, 99)

    # Summary box table
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_meta.autofit = False

    meta_items = [
        ("Project Name", "Agrirent - Agricultural Equipment Rental System"),
        ("Technology Stack", "Next.js 15 (App Router), TypeScript, Supabase PostgreSQL, Tailwind CSS"),
        ("Authentication & RBAC", "Supabase Auth with Role-Based Access Control (Admin, Owner, Customer)"),
        ("Key Architecture", "Server Actions, Row Level Security (RLS), Cloud Storage, Responsive Base UI"),
        ("Document Version", "Version 2.0 (Complete Feature Specification with Screenshots)")
    ]

    for i, (k, v) in enumerate(meta_items):
        row = table_meta.rows[i]
        c0 = row.cells[0]
        c1 = row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.2)
        
        c0.text = k
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(10)
        c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(22, 101, 52)
        
        c1.text = v
        c1.paragraphs[0].runs[0].font.size = Pt(10)
        
        set_cell_background(c0, "F0FDF4")
        set_cell_background(c1, "F9FAFB")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)

    doc.add_page_break()

    # ==========================================
    # 2. EXECUTIVE SUMMARY & ARCHITECTURE
    # ==========================================
    add_heading_styled(doc, "1. Executive Summary & Overview", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Agrirent is a modern, full-stack digital marketplace created to solve agricultural mechanization challenges "
        "by connecting local equipment owners (tractors, harvesters, tillers, rotavators, sprayers) directly with farmers and customers. "
        "The platform eliminates middle-men, offers transparent hourly/daily pricing, prevents double-booking through interactive availability calendars, "
        "and incorporates end-to-end identity verification and administrator moderation."
    )

    add_heading_styled(doc, "System Architecture", level=2)
    add_bullet_point(doc, "Frontend Framework", "Next.js 15 (App Router) leveraging React Server Components (RSC) for lightning-fast initial load times and client components for real-time reactivity.")
    add_bullet_point(doc, "Database & Auth Engine", "Supabase PostgreSQL with Row Level Security (RLS) policies, programmatic auto-confirmation, and automated PostgreSQL triggers.")
    add_bullet_point(doc, "Cloud Document & Media Storage", "Supabase Storage buckets for high-resolution equipment listings, customer avatars, and Aadhaar identity verification files.")
    add_bullet_point(doc, "Role-Based Access Control (RBAC)", "Strict multi-tenant privilege separation dividing users into Customer, Equipment Owner, and Platform Administrator tiers.")
    add_bullet_point(doc, "Design System", "Tailwind CSS paired with modern Base UI primitives, Lucide icons, glassmorphic headers, and dedicated mobile viewports with bottom navigation.")

    # ==========================================
    # 3. FEATURE: LANDING PAGE & MARKETPLACE
    # ==========================================
    add_heading_styled(doc, "2. Landing Page & Equipment Marketplace", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The Agrirent landing page provides an inviting, high-converting entry point for farmers and equipment owners. "
        "It features dynamic hero search by category, location, and price, along with direct highlights of top-rated machinery."
    )

    add_image_safe(doc, os.path.join(img_dir, '01_landing_page.png'), "Agrirent Landing Page Hero & Modern Header", width=Inches(6.2))

    add_heading_styled(doc, "Marketplace & Filter Engine", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The Equipment Marketplace (/equipment) lets farmers browse through tractors, harvesters, and attachments. "
        "Key capabilities include:"
    )
    add_bullet_point(doc, "Real-Time Multi-Filter", "Filter by machinery category, daily rental rate range, owner rating, and geolocation.")
    add_bullet_point(doc, "Instant Booking Modal", "Direct booking dialog displaying equipment specifications, rental terms, and total price calculation.")
    add_bullet_point(doc, "Owner Self-Management Integration", "When an equipment owner visits their own listing, the system automatically detects ownership and displays a 'Manage & Edit Equipment' button instead of 'Request to Rent'.")

    add_image_safe(doc, os.path.join(img_dir, '02_equipment_marketplace.png'), "Equipment Marketplace with Search & Live Categories", width=Inches(6.2))

    # ==========================================
    # 4. FEATURE: AUTHENTICATION & AADHAAR VERIFICATION
    # ==========================================
    add_heading_styled(doc, "3. Authentication & Owner Aadhaar Verification", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Agrirent features a robust authentication flow supporting dual registration paths for Customers and Equipment Owners. "
        "All user registrations capture Full Name, Email, Phone Number, and Secure Passwords with server-side validation."
    )

    add_heading_styled(doc, "Customer vs. Owner Signup Tabs", level=2)
    add_bullet_point(doc, "Customer Registration", "Instant onboarding allowing users to explore equipment, save favorites, and place booking requests.")
    add_bullet_point(doc, "Owner Registration with Aadhaar Upload", "Equipment owners are required to upload an official Aadhaar Card / Identity Proof document (.jpg, .jpeg, .png, .pdf) up to 10MB during registration.")
    add_bullet_point(doc, "Automated Profile & Phone Sync", "PostgreSQL database triggers sync the phone number and metadata into the public profiles table upon registration.")

    add_image_safe(doc, os.path.join(img_dir, '03_signup_page.png'), "Customer Registration Form Tab", width=Inches(5.0))
    add_image_safe(doc, os.path.join(img_dir, '05_owner_signup_aadhaar.png'), "Owner Registration with Integrated Aadhaar Document Upload", width=Inches(5.0))
    add_image_safe(doc, os.path.join(img_dir, '04_login_page.png'), "Secure Sign-in & Authentication Portal", width=Inches(5.0))

    # ==========================================
    # 5. FEATURE: ADMIN DASHBOARD & OPERATIONS
    # ==========================================
    add_heading_styled(doc, "4. Administrator Dashboard & Platform Controls", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The Agrirent Admin Dashboard (/dashboard/admin) provides centralized governance, moderation, and analytics across all platform operations. "
        "It features live KPI metrics, prominent action items, and dedicated management interfaces."
    )

    add_heading_styled(doc, "Admin Features & Visual Red Logout Button", level=2)
    add_bullet_point(doc, "Platform Operations Banner", "One-click shortcuts to Verifications, Equipment Approvals, and Financial Transactions.")
    add_bullet_point(doc, "Real-Time Financial & Booking KPIs", "Color-coded cards displaying Platform Revenue (INR), Total Bookings, Total Users, and Pending Approvals compared against previous day metrics.")
    add_bullet_point(doc, "Red Logout Button", "Prominent red logout button with high-contrast icon for quick, secure administrative session termination.")

    add_image_safe(doc, os.path.join(img_dir, '07_admin_dashboard.png'), "Administrator Dashboard with KPI Analytics and Operations Center", width=Inches(6.2))

    add_heading_styled(doc, "Owner Verifications & Aadhaar Document Inspector", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "On the Admin Owner Verifications page (/dashboard/admin/owners), administrators inspect owner applicants and verify their credentials before granting listing privileges:"
    )
    add_bullet_point(doc, "Direct Aadhaar Viewer", "Clicking 'View Aadhaar Card ↗' opens the uploaded Aadhaar document in full resolution.")
    add_bullet_point(doc, "Approve & Reject Workflow", "Admins can approve or reject applications in one click. Approving automatically grants the owner role and unlocks fleet management.")

    add_image_safe(doc, os.path.join(img_dir, '08_admin_owner_verifications.png'), "Owner Verifications Screen with Direct Aadhaar Inspection", width=Inches(6.2))

    add_heading_styled(doc, "User Management & Admin Edit User Dialog", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "The User Management panel (/dashboard/admin/users) provides total visibility over all registered accounts. "
        "Admins have full capability to edit user details directly through a Base UI interactive modal:"
    )
    add_bullet_point(doc, "Edit Contact Information", "Admins can update email addresses and phone numbers when users request phone or email changes.")
    add_bullet_point(doc, "Role Management", "Admins can seamlessly switch account roles between Customer, Equipment Owner, and Admin.")
    add_bullet_point(doc, "Dual Table & Card Representation", "Full tabular presentation on desktop and compact card view on mobile.")

    # ==========================================
    # 6. FEATURE: OWNER FLEET & BOOKING MANAGEMENT
    # ==========================================
    add_heading_styled(doc, "5. Equipment Owner Fleet & Booking Management", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Verified equipment owners have access to the Owner Dashboard (/dashboard/owner) to manage equipment listings, "
        "track incoming booking requests, and visualize booking schedules."
    )

    add_heading_styled(doc, "Fleet Management & Booking Calendar", level=2)
    add_bullet_point(doc, "Add & Edit Machinery", "Upload multiple photos, set daily/hourly rates, specify operating location, and define equipment technical specs.")
    add_bullet_point(doc, "Optimized Booking Calendar", "Compact, responsive calendar view allowing owners to see reserved vs. available dates at a glance on both desktop and mobile.")
    add_bullet_point(doc, "Booking Approval Flow", "Owners can accept or reject rental requests with live status updates and automatic SMS/Email notifications.")

    add_image_safe(doc, os.path.join(img_dir, '09_owner_dashboard_calendar.png'), "Owner Dashboard with Booking Schedule and Fleet Summary", width=Inches(6.2))
    add_image_safe(doc, os.path.join(img_dir, '10_owner_mobile_fleet.png'), "Owner Fleet Management on Mobile Viewport", width=Inches(5.0))

    # ==========================================
    # 7. FEATURE: MOBILE OPTIMIZATION & NAVIGATION
    # ==========================================
    add_heading_styled(doc, "6. Mobile Responsive Design & Role-Aware Navigation", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Because agricultural machinery owners and farmers frequently access the platform from smartphones in the field, "
        "Agrirent is built with a mobile-first philosophy."
    )

    add_bullet_point(doc, "Role-Aware Bottom Navigation Bar", "Fixed bottom navigation bar automatically alters its tabs based on the logged-in role (Explore/Bookings/Favorites for Customers; Dashboard/My Fleet/Bookings/Profile for Owners).")
    add_bullet_point(doc, "Smart Top-Level Page Routers", "Dedicated server-side handlers for /profile, /bookings, and /favorites ensure 100% reliable navigation without 404 errors.")
    add_bullet_point(doc, "Responsive Data Cards", "Complex desktop tables on Admin, Owner, and User pages dynamically transform into clean mobile cards with action buttons and badge indicators.")

    add_image_safe(doc, os.path.join(img_dir, '06_mobile_equipment_marketplace.png'), "Mobile Marketplace View with Quick Category Filtering", width=Inches(4.5))
    add_image_safe(doc, os.path.join(img_dir, '11_admin_mobile_dashboard.png'), "Admin Dashboard in Mobile View", width=Inches(4.5))

    # ==========================================
    # 8. DATABASE SCHEMA & SECURITY
    # ==========================================
    add_heading_styled(doc, "7. Database Schema & Security Policies", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Agrirent is backed by PostgreSQL tables in Supabase with strict foreign keys, cascade rules, and Row Level Security (RLS):"
    )

    schema_table = doc.add_table(rows=7, cols=3)
    schema_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    schema_table.autofit = False

    schema_rows = [
        ("Table Name", "Key Columns", "Description & Security"),
        ("profiles", "id, full_name, email, phone, role, avatar_url", "Stores user profile data synced with auth.users. Read public, update owner/admin."),
        ("equipment", "id, owner_id, title, category_id, daily_rate, location, status", "Agricultural machinery listings. Public read approved, owners manage own."),
        ("bookings", "id, equipment_id, customer_id, start_date, end_date, total_price, status", "Rental reservations. Customers manage own; equipment owners manage received requests."),
        ("owner_requests", "id, user_id, business_name, identity_document_url, status", "Owner verification records storing Aadhaar card URL. Admin review & update."),
        ("payments", "id, booking_id, amount, payment_method, status", "Transaction records storing platform commission and owner payouts."),
        ("reviews", "id, equipment_id, user_id, rating, comment", "Customer feedback and ratings displayed on machinery cards.")
    ]

    for i, (col1, col2, col3) in enumerate(schema_rows):
        row = schema_table.rows[i]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
        c0.width = Inches(1.3)
        c1.width = Inches(2.3)
        c2.width = Inches(2.8)
        
        c0.text = col1
        c1.text = col2
        c2.text = col3
        
        if i == 0:
            for cell in (c0, c1, c2):
                set_cell_background(cell, "15803D") # Green header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_background(c0, "F0FDF4" if i%2==1 else "FFFFFF")
            set_cell_background(c1, "F9FAFB" if i%2==1 else "FFFFFF")
            set_cell_background(c2, "F9FAFB" if i%2==1 else "FFFFFF")
            c0.paragraphs[0].runs[0].font.bold = True
            c0.paragraphs[0].runs[0].font.size = Pt(9.5)
            c1.paragraphs[0].runs[0].font.size = Pt(9.5)
            c2.paragraphs[0].runs[0].font.size = Pt(9.5)
            
        for cell in (c0, c1, c2):
            set_cell_margins(cell, 50, 50, 80, 80)

    # ==========================================
    # 9. CONCLUSION & DEPLOYMENT
    # ==========================================
    add_heading_styled(doc, "8. Conclusion & Git Repository Summary", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.add_run(
        "Agrirent delivers an enterprise-grade, high-performance web platform tailored for the modern agricultural ecosystem. "
        "With full role-based permissions, automated Aadhaar document verification, intuitive booking calendars, and mobile-first optimization, "
        "the application ensures security, scalability, and ease of use for all stakeholders."
    )

    output_path = os.path.join(base_dir, 'Agrirent_Project_Documentation.docx')
    doc.save(output_path)
    print(f"Successfully generated documentation at: {output_path}")

if __name__ == '__main__':
    generate_document()
